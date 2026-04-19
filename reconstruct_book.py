#!/usr/bin/env python3
"""
OCR Book Reconstruction Agent
==============================
Processes book page images, detects page numbers, renames files,
extracts text via OCR, and reconstructs a formatted DOCX document.
"""

import os
import re
import sys
import shutil
import logging
from collections import Counter, defaultdict
from pathlib import Path

import cv2
import numpy as np
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
IMAGE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "images")
OUTPUT_DOCX = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reconstructed_book.docx")
LOG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "processing_log.txt")
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
MAX_PAGE_NUMBER = 200  # Reasonable upper bound for page numbers

# Arabic-Indic numeral mapping  ٠١٢٣٤٥٦٧٨٩
ARABIC_INDIC_MAP = {
    "\u0660": "0", "\u0661": "1", "\u0662": "2", "\u0663": "3", "\u0664": "4",
    "\u0665": "5", "\u0666": "6", "\u0667": "7", "\u0668": "8", "\u0669": "9",
}

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(LOG_FILE, mode="w", encoding="utf-8"),
    ],
)
log = logging.getLogger(__name__)

# =========================================================================
# Step 1 – Image Discovery
# =========================================================================

def discover_images(root_dir: str) -> list[str]:
    """Recursively find all image files."""
    images = []
    for dirpath, _, filenames in os.walk(root_dir):
        for fname in filenames:
            if Path(fname).suffix.lower() in IMAGE_EXTENSIONS:
                images.append(os.path.join(dirpath, fname))
    images.sort()
    log.info("Step 1: Discovered %d images in %s", len(images), root_dir)
    return images

# =========================================================================
# Step 2 – Page Number Detection
# =========================================================================

def _convert_arabic_indic(text: str) -> str:
    for ar, en in ARABIC_INDIC_MAP.items():
        text = text.replace(ar, en)
    return text


def _preprocess_for_page_number(img: Image.Image) -> Image.Image:
    """Enhance a small crop for page-number OCR."""
    img = img.convert("L")
    img = ImageEnhance.Contrast(img).enhance(3.0)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    # Binarise
    arr = np.array(img)
    _, binarised = cv2.threshold(arr, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return Image.fromarray(binarised)


def _detect_numbers_in_crop(crop: Image.Image) -> list[int]:
    """Run OCR on a small crop and extract integers."""
    processed = _preprocess_for_page_number(crop)
    found: list[int] = []
    for lang in ("ara", "eng", "ara+eng"):
        for psm in (7, 13, 6):
            try:
                raw = pytesseract.image_to_string(processed, lang=lang, config=f"--psm {psm}")
                raw = _convert_arabic_indic(raw)
                nums = re.findall(r"\b(\d{1,3})\b", raw)
                for n_str in nums:
                    n = int(n_str)
                    if 1 <= n <= MAX_PAGE_NUMBER:
                        found.append(n)
            except Exception:
                pass
    return found


def detect_page_number(img_path: str) -> tuple[int | None, str]:
    """
    Detect page number from the top region of an image.

    Returns (page_number_or_None, confidence_level).
    confidence_level is one of: 'high', 'medium', 'low', 'uncertain'.
    """
    img = Image.open(img_path)
    w, h = img.size

    # Define focused corner/edge regions where page numbers typically appear
    margin_h = int(h * 0.07)  # top 7 %
    margin_w = int(w * 0.20)  # side 20 %

    regions = {
        "top_left":   (0, 0, margin_w, margin_h),
        "top_center": (margin_w, 0, w - margin_w, margin_h),
        "top_right":  (w - margin_w, 0, w, margin_h),
    }

    all_candidates: list[int] = []
    region_candidates: dict[str, list[int]] = {}

    for name, bbox in regions.items():
        crop = img.crop(bbox)
        nums = _detect_numbers_in_crop(crop)
        region_candidates[name] = nums
        all_candidates.extend(nums)

    if not all_candidates:
        return None, "uncertain"

    counter = Counter(all_candidates)
    best_num, best_count = counter.most_common(1)[0]

    # Confidence scoring
    total_detections = len(all_candidates)
    agreement_ratio = best_count / total_detections

    if best_count >= 6 and agreement_ratio >= 0.5:
        confidence = "high"
    elif best_count >= 3 and agreement_ratio >= 0.35:
        confidence = "medium"
    elif best_count >= 2:
        confidence = "low"
    else:
        confidence = "uncertain"

    return best_num, confidence

# =========================================================================
# Step 3 – Rename Files
# =========================================================================

def rename_files(
    image_paths: list[str],
    detections: dict[str, tuple[int | None, str]],
) -> dict[str, str]:
    """
    Rename images based on detected page numbers.
    Returns mapping old_path -> new_path.
    """
    rename_map: dict[str, str] = {}
    used_names: set[str] = set()

    for path in image_paths:
        page_num, confidence = detections[path]
        directory = os.path.dirname(path)
        ext = Path(path).suffix.lower()

        if page_num is not None and confidence in ("high", "medium", "low"):
            base = f"page_{page_num}"
            new_name = f"{base}{ext}"
            # Handle duplicates
            counter = 1
            while new_name in used_names:
                new_name = f"{base}_dup{counter}{ext}"
                counter += 1
                log.warning("Duplicate page number %d detected, using %s", page_num, new_name)
        else:
            stem = Path(path).stem
            new_name = f"{stem}_unknown{ext}"
            dup_counter = 1
            while new_name in used_names:
                new_name = f"{stem}_unknown_{dup_counter}{ext}"
                dup_counter += 1

        used_names.add(new_name)
        new_path = os.path.join(directory, new_name)
        rename_map[path] = new_path

    # Perform renames (use intermediate names to avoid conflicts)
    temp_map: dict[str, str] = {}
    for old, new in rename_map.items():
        if old != new:
            tmp = old + ".tmp_rename"
            shutil.move(old, tmp)
            temp_map[tmp] = new

    for tmp, new in temp_map.items():
        shutil.move(tmp, new)
        log.info("Renamed: %s -> %s", os.path.basename(tmp).replace(".tmp_rename", ""), os.path.basename(new))

    return rename_map

# =========================================================================
# Step 4 – Ordering Strategy
# =========================================================================

def order_pages(
    rename_map: dict[str, str],
    detections: dict[str, tuple[int | None, str]],
) -> list[str]:
    """
    Sort pages: valid page numbers first (ascending), then unknowns by
    original filename order.
    """
    numbered: list[tuple[int, str]] = []
    unknowns: list[tuple[str, str]] = []

    for old_path, new_path in rename_map.items():
        page_num, confidence = detections[old_path]
        if page_num is not None and confidence in ("high", "medium", "low"):
            numbered.append((page_num, new_path))
        else:
            unknowns.append((old_path, new_path))

    # Sort numbered pages
    numbered.sort(key=lambda x: x[0])

    # Sort unknowns by original filename
    unknowns.sort(key=lambda x: os.path.basename(x[0]))

    ordered = [p for _, p in numbered] + [p for _, p in unknowns]

    # Log ordering
    log.info("Step 4: Page ordering (total %d):", len(ordered))
    for i, p in enumerate(ordered, 1):
        log.info("  %3d. %s", i, os.path.basename(p))

    # Check for gaps in page numbers
    page_nums = [n for n, _ in numbered]
    if page_nums:
        expected = set(range(min(page_nums), max(page_nums) + 1))
        missing = expected - set(page_nums)
        if missing:
            log.warning("Missing page numbers in sequence: %s", sorted(missing))

    return ordered

# =========================================================================
# Step 5 – OCR Text Extraction
# =========================================================================

def _preprocess_for_ocr(img_path: str) -> Image.Image:
    """Preprocess image for high-accuracy OCR."""
    img = cv2.imread(img_path)
    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # Denoise
    denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
    # Adaptive threshold
    binary = cv2.adaptiveThreshold(
        denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    return Image.fromarray(binary)


def _fix_common_ocr_errors(text: str) -> str:
    """Fix common OCR mis-reads."""
    # Common substitution errors
    replacements = {
        "ﻻ": "لا",
        "ﻷ": "لأ",
        "ﻹ": "لإ",
        "ﻵ": "لآ",
    }
    for wrong, right in replacements.items():
        text = text.replace(wrong, right)

    # Remove isolated single characters that are likely noise
    # But preserve Arabic single-character words (و، ف، ب، etc.)
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if line:
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


def extract_text(img_path: str) -> str:
    """Extract text from an image using OCR."""
    try:
        processed = _preprocess_for_ocr(img_path)

        # Try Arabic first, then combined
        text_ara = pytesseract.image_to_string(processed, lang="ara", config="--psm 6")
        text_combined = pytesseract.image_to_string(processed, lang="ara+eng", config="--psm 6")

        # Use the longer text (more content detected)
        text = text_ara if len(text_ara) >= len(text_combined) else text_combined
        text = _convert_arabic_indic(text)
        text = _fix_common_ocr_errors(text)

        return text.strip()
    except Exception as e:
        log.error("OCR failed for %s: %s", img_path, e)
        return ""

# =========================================================================
# Step 6 – Layout & Structure Detection
# =========================================================================

def _classify_line(line: str, avg_line_length: float) -> str:
    """
    Classify a line as title, heading, or paragraph based on heuristics.
    """
    stripped = line.strip()
    if not stripped:
        return "empty"

    # Very short centered-looking text might be a title/heading
    is_short = len(stripped) < avg_line_length * 0.4
    has_colon = stripped.endswith(":")
    starts_with_number = bool(re.match(r"^\d+[\.\)\-]", stripped))
    starts_with_bullet = stripped.startswith(("-", "•", "–", "*", "◆"))

    if is_short and len(stripped) < 30:
        return "title"
    elif is_short and len(stripped) < 60:
        return "heading"
    elif starts_with_bullet or starts_with_number:
        return "list_item"
    elif has_colon and is_short:
        return "heading"
    else:
        return "paragraph"


def detect_structure(text: str) -> list[dict]:
    """
    Analyze text and return structured elements.
    Each element: {'type': 'title'|'heading'|'paragraph'|'list_item', 'text': str}
    """
    lines = text.split("\n")
    non_empty = [l for l in lines if l.strip()]
    if not non_empty:
        return []

    avg_len = sum(len(l.strip()) for l in non_empty) / len(non_empty)

    elements: list[dict] = []
    current_paragraph: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_paragraph:
                elements.append({
                    "type": "paragraph",
                    "text": " ".join(current_paragraph),
                })
                current_paragraph = []
            continue

        line_type = _classify_line(stripped, avg_len)

        if line_type in ("title", "heading", "list_item"):
            # Flush current paragraph first
            if current_paragraph:
                elements.append({
                    "type": "paragraph",
                    "text": " ".join(current_paragraph),
                })
                current_paragraph = []
            elements.append({"type": line_type, "text": stripped})
        else:
            current_paragraph.append(stripped)

    # Flush remaining
    if current_paragraph:
        elements.append({
            "type": "paragraph",
            "text": " ".join(current_paragraph),
        })

    return elements

# =========================================================================
# Step 7 – Document Reconstruction
# =========================================================================

def _fix_broken_sentences(pages_text: list[str]) -> list[str]:
    """
    Fix sentences that are broken across page boundaries.
    """
    if len(pages_text) <= 1:
        return pages_text

    fixed: list[str] = []
    for i, text in enumerate(pages_text):
        if i == 0:
            fixed.append(text)
            continue

        prev = fixed[-1]
        current = text

        # Check if previous page ends mid-sentence (no period/question mark)
        prev_stripped = prev.rstrip()
        ends_with_punct = prev_stripped and prev_stripped[-1] in ".!?؟。،"

        if not ends_with_punct and current:
            # Find first line of current
            first_line_end = current.find("\n")
            if first_line_end == -1:
                first_line = current
                rest = ""
            else:
                first_line = current[:first_line_end]
                rest = current[first_line_end + 1:]

            # Merge last line of previous with first line of current
            fixed[-1] = prev_stripped + " " + first_line.strip()
            if rest:
                fixed.append(rest)
            else:
                fixed.append("")
        else:
            fixed.append(current)

    return fixed


def reconstruct_document(
    ordered_paths: list[str],
) -> list[tuple[str, list[dict]]]:
    """
    Extract text from all pages and reconstruct into structured elements.
    Returns list of (page_label, elements) tuples.
    """
    raw_texts: list[str] = []
    page_labels: list[str] = []

    for path in ordered_paths:
        log.info("Step 5: Extracting text from %s", os.path.basename(path))
        text = extract_text(path)
        raw_texts.append(text)
        page_labels.append(os.path.basename(path))

    # Fix broken sentences across pages
    log.info("Step 7: Fixing broken sentences across %d pages", len(raw_texts))
    fixed_texts = _fix_broken_sentences(raw_texts)

    # Detect structure for each page
    result: list[tuple[str, list[dict]]] = []
    for label, text in zip(page_labels, fixed_texts):
        elements = detect_structure(text)
        result.append((label, elements))
        log.info("Step 6: %s -> %d elements detected", label, len(elements))

    return result

# =========================================================================
# Step 8 & 9 – DOCX Formatting & Output
# =========================================================================

def _setup_document_styles(doc: Document) -> None:
    """Configure document styles for Arabic RTL text."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Traditional Arabic"
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # RTL

    for level, size in [("Heading 1", 24), ("Heading 2", 20), ("Heading 3", 16)]:
        if level in doc.styles:
            h_style = doc.styles[level]
            h_style.font.name = "Traditional Arabic"
            h_style.font.size = Pt(size)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)


def generate_docx(
    pages_data: list[tuple[str, list[dict]]],
    output_path: str,
) -> None:
    """Generate the final DOCX file."""
    doc = Document()
    _setup_document_styles(doc)

    # Track for quality control
    total_elements = 0
    uncertain_count = 0
    seen_texts: set[str] = set()

    for page_idx, (page_label, elements) in enumerate(pages_data):
        if page_idx > 0:
            # Add a subtle page break between pages
            doc.add_page_break()

        for elem in elements:
            text = elem["text"].strip()
            if not text:
                continue

            # Quality control: deduplicate
            text_hash = text[:100]  # Use first 100 chars as key
            if text_hash in seen_texts and len(text) > 20:
                log.warning("Duplicate text detected (skipping): %.50s...", text)
                continue
            seen_texts.add(text_hash)

            # Quality control: mark uncertain OCR
            # (text with too many special chars or very short garbled text)
            garble_ratio = sum(1 for c in text if c in "□■◻◼▪▫●○◆◇★☆") / max(len(text), 1)
            if garble_ratio > 0.3:
                text = f"[UNCERTAIN] {text}"
                uncertain_count += 1

            elem_type = elem["type"]
            if elem_type == "title":
                p = doc.add_heading(text, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif elem_type == "heading":
                p = doc.add_heading(text, level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif elem_type == "list_item":
                p = doc.add_paragraph(text, style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            total_elements += 1

    doc.save(output_path)
    log.info("Step 9: Generated %s (%d elements, %d uncertain)", output_path, total_elements, uncertain_count)

# =========================================================================
# Step 10 – Quality Control
# =========================================================================

def quality_control(
    ordered_paths: list[str],
    detections: dict[str, tuple[int | None, str]],
    rename_map: dict[str, str],
) -> None:
    """Run quality checks and log results."""
    log.info("=" * 60)
    log.info("QUALITY CONTROL REPORT")
    log.info("=" * 60)

    # Check page ordering
    page_nums = []
    for old_path in sorted(detections.keys()):
        num, conf = detections[old_path]
        if num is not None and conf in ("high", "medium", "low"):
            page_nums.append(num)

    sorted_nums = sorted(set(page_nums))
    if sorted_nums:
        log.info("Detected page range: %d - %d", min(sorted_nums), max(sorted_nums))
        expected = set(range(min(sorted_nums), max(sorted_nums) + 1))
        missing = expected - set(sorted_nums)
        if missing:
            log.warning("Missing pages: %s", sorted(missing))
        else:
            log.info("No missing pages in detected range")

    # Duplicate page numbers
    dup_nums = [n for n, c in Counter(page_nums).items() if c > 1]
    if dup_nums:
        log.warning("Duplicate page numbers detected: %s", dup_nums)

    # Uncertain detections
    uncertain = [(p, d) for p, d in detections.items() if d[1] == "uncertain"]
    if uncertain:
        log.warning("%d images with uncertain page numbers:", len(uncertain))
        for p, d in uncertain:
            log.warning("  %s", os.path.basename(p))

    # Summary
    high = sum(1 for _, (_, c) in detections.items() if c == "high")
    medium = sum(1 for _, (_, c) in detections.items() if c == "medium")
    low = sum(1 for _, (_, c) in detections.items() if c == "low")
    unc = sum(1 for _, (_, c) in detections.items() if c == "uncertain")

    log.info("Detection confidence: high=%d, medium=%d, low=%d, uncertain=%d",
             high, medium, low, unc)
    log.info("Total pages processed: %d", len(ordered_paths))
    log.info("=" * 60)

# =========================================================================
# Main Pipeline
# =========================================================================

def main():
    log.info("=" * 60)
    log.info("OCR Book Reconstruction Agent – Starting")
    log.info("=" * 60)

    # Step 1: Image Discovery
    image_paths = discover_images(IMAGE_DIR)
    if not image_paths:
        log.error("No images found in %s", IMAGE_DIR)
        sys.exit(1)

    # Step 2: Page Number Detection
    log.info("Step 2: Detecting page numbers...")
    detections: dict[str, tuple[int | None, str]] = {}
    for path in image_paths:
        page_num, confidence = detect_page_number(path)
        detections[path] = (page_num, confidence)
        log.info("  %s -> page=%s confidence=%s",
                 os.path.basename(path), page_num, confidence)

    # Step 3: Rename Files
    log.info("Step 3: Renaming files...")
    rename_map = rename_files(image_paths, detections)

    # Step 4: Ordering
    log.info("Step 4: Determining page order...")
    ordered_paths = order_pages(rename_map, detections)

    # Steps 5-7: Extract, Structure, Reconstruct
    log.info("Steps 5-7: OCR extraction and reconstruction...")
    pages_data = reconstruct_document(ordered_paths)

    # Steps 8-9: DOCX generation
    log.info("Steps 8-9: Generating DOCX...")
    generate_docx(pages_data, OUTPUT_DOCX)

    # Step 10: Quality Control
    quality_control(ordered_paths, detections, rename_map)

    log.info("Done! Output: %s", OUTPUT_DOCX)


if __name__ == "__main__":
    main()
